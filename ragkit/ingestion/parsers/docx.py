"""DOCX parser implementation."""

from __future__ import annotations

import io
import os
import shutil
import subprocess
import tempfile

import structlog

from ragkit.config.schema import ParsingConfig
from ragkit.ingestion.parsers.base import BaseParser, ParsedDocument
from ragkit.ingestion.sources.base import RawDocument


def _extract_with_unstructured(raw_bytes: bytes, file_type: str = "docx") -> str | None:
    try:
        if file_type == "doc":
            from unstructured.partition.doc import partition_doc

            elements = partition_doc(file=io.BytesIO(raw_bytes))
        else:
            from unstructured.partition.docx import partition_docx

            elements = partition_docx(file=io.BytesIO(raw_bytes))
    except Exception:
        return None

    parts: list[str] = []
    for element in elements:
        text = getattr(element, "text", None)
        if text:
            parts.append(text)
        else:
            parts.append(str(element))
    return "\n".join(parts).strip()


def _extract_with_antiword(raw_bytes: bytes) -> str | None:
    if not shutil.which("antiword"):
        return None
    tmp_name: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(raw_bytes)
            tmp.flush()
            tmp_name = tmp.name
        result = subprocess.run(
            ["antiword", tmp_name],
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout.decode("utf-8", errors="replace").strip()
    except Exception:
        pass
    finally:
        if tmp_name:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
    return None


def _extract_with_python_docx(raw_bytes: bytes) -> str | None:
    try:
        import docx
    except Exception:
        return None

    doc = docx.Document(io.BytesIO(raw_bytes))
    parts = [para.text for para in doc.paragraphs if para.text]
    return "\n".join(parts).strip()


def _fallback_decode(raw_bytes: bytes) -> str:
    return raw_bytes.decode("utf-8", errors="ignore") or raw_bytes.decode(
        "latin-1", errors="ignore"
    )


class DOCXParser(BaseParser):
    def __init__(self, config: ParsingConfig | None = None):
        self.config = config or ParsingConfig()
        self.logger = structlog.get_logger()

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in {"docx", "doc"}

    async def parse(self, raw_doc: RawDocument) -> ParsedDocument:
        content = raw_doc.content
        file_type = raw_doc.file_type.lower() if raw_doc.file_type else "docx"
        if isinstance(content, str):
            text = content
        elif file_type == "doc":
            if not _has_doc_tools():
                self.logger.warning(
                    "doc_parser_missing_dependencies",
                    message="antiword and soffice not found, .doc extraction may produce garbled text",
                    source=raw_doc.source_path,
                )
            text = (
                _extract_with_unstructured(content, file_type="doc")
                or _extract_with_antiword(content)
                or _fallback_decode(content)
            )
        else:
            text = (
                _extract_with_unstructured(content, file_type="docx")
                or _extract_with_python_docx(content)
                or _fallback_decode(content)
            )

        metadata = dict(raw_doc.metadata)
        metadata.setdefault("file_type", file_type)
        return ParsedDocument(content=text, metadata=metadata, structure=None)


def _has_doc_tools() -> bool:
    return bool(shutil.which("antiword") or shutil.which("soffice"))
